import streamlit as st
import anthropic
import fitz  # PyMuPDF
from docx import Document
from io import BytesIO
import time


st.set_page_config(page_title="TraductorProAI", page_icon="🌐", layout="wide")


st.markdown("""
    <style>
    .main { background-color: #f8f9fa; }
    .stButton>button { width: 100%; border-radius: 8px; background-color: #ff4b4b; color: white; height: 3em; }
    </style>
    """, unsafe_allow_html=True)


client = None
if "CLAUDE_API_KEY" in st.secrets:
    api_key = st.secrets["CLAUDE_API_KEY"]
    client = anthropic.Anthropic(api_key=api_key)
else:
    with st.sidebar:
        st.title("🔐 Seguridad")
        api_key = st.text_input("Introduce tu Anthropic API Key", type="password")
        if api_key:
            client = anthropic.Anthropic(api_key=api_key)



def extraer_texto_pdf(archivo):
    doc = fitz.open(stream=archivo.read(), filetype="pdf")
    texto = ""
    for pagina in doc:
        texto += pagina.get_text()
    return texto

def dividir_texto(texto, limite_palabras=1200):
    """Divide el texto en fragmentos más pequeños para mejor precisión de Claude"""
    palabras = texto.split()
    for i in range(0, len(palabras), limite_palabras):
        yield " ".join(palabras[i:i + limite_palabras])

def traducir_con_claude(bloque, idioma_destino):
    try:
        
        message = client.messages.create(
            model="claude-3-5-sonnet-20240620",
            max_tokens=4000,
            temperature=0,
            system=f"Eres un traductor de élite. Traduce el siguiente texto al {idioma_destino}. "
                   f"Mantén un tono profesional y respeta la estructura de párrafos. "
                   f"Solo entrega el texto traducido, sin introducciones ni comentarios.",
            messages=[{"role": "user", "content": bloque}]
        )
        return message.content[0].text
    except Exception as e:
        return f"\n[Error en traducción: {str(e)}]\n"

def crear_docx(texto_traducido):
    doc = Document()
    doc.add_heading('Traducción Generada por TraductorProAI', 0)
    for parrafo in texto_traducido.split('\n'):
        if parrafo.strip():
            doc.add_paragraph(parrafo)
    
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


st.title("🌐 TraductorProAI")
st.markdown("### El poder de Claude 3.5 para tus documentos PDF")
st.divider()

if not client:
    st.warning("⚠️ Esperando API Key para activar los motores de IA...")
else:
    col1, col2 = st.columns([2, 1])
    
    with col1:
        archivo_pdf = st.file_uploader("Carga tu documento PDF", type=["pdf"])
    
    with col2:
        idioma = st.selectbox("Idioma objetivo", ["Inglés", "Español", "Francés", "Alemán", "Italiano", "Portugués", "Chino", "Japonés"])

    if archivo_pdf and st.button("🚀 Iniciar Traducción Profesional"):
        texto_original = extraer_texto_pdf(archivo_pdf)
        fragmentos = list(dividir_texto(texto_original))
        
        st.info(f"Documento analizado. Se procesará en {len(fragmentos)} bloques de alta fidelidad.")
        
        traduccion_acumulada = ""
        progreso = st.progress(0)
        
        for idx, bloque in enumerate(fragmentos):
            with st.spinner(f"Traduciendo bloque {idx+1} de {len(fragmentos)}..."):
                resultado = traducir_con_claude(bloque, idioma)
                traduccion_acumulada += resultado + "\n\n"
                
                
                progreso.progress((idx + 1) / len(fragmentos))
                time.sleep(0.5) 

        st.success("✅ ¡Traducción finalizada con éxito!")
        
      
        with st.expander("Ver previsualización del texto"):
            st.write(traduccion_acumulada)
        
        
        doc_final = crear_docx(traduccion_acumulada)
        
        st.download_button(
            label="📥 Descargar Documento Word (.docx)",
            data=doc_final,
            file_name=f"Traduccion_{idioma}_TraductorProAI.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
